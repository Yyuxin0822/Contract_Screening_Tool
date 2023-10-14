from flask import Flask, request, redirect, url_for, render_template, send_file
from waitress import serve
from form import UploadFileForm
from werkzeug.utils import secure_filename
import docx
from docx import Document
from contract import screening
import os

app=Flask(__name__)
app.config['SECRET_KEY']="secretkey"
app.config['UPLOAD_FOLDER'] = 'static/uploads'

notice={0:""}
doc={0:Document()}

@app.route("/", methods=["GET", "POST"])
def index(id=0):
    form=UploadFileForm()
    if form.validate_on_submit():
        if len(doc)>1:
            del doc[1]
            del notice[1]
        file=form.file.data
        id=1
        #validate file type to be .docx
        if file.filename.endswith('.docx'):
            #filepath=os.path.join(os.path.abspath(os.path.dirname(__file__)),app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            #file.save(filepath)
            # doc_dir="./static/uploads/_template.docx"
            doc[id]= docx.Document(file)
            notice[id]=f"Congrats! File Executed Successfully!"
        else:
            notice[id]=f"File must end with .docx"     
        return redirect(url_for('index',id=id))
    id=len(notice)-1
    return render_template("index.html",form=form,notice=notice[id])



@app.route("/template")
def template():
    doc_dir="./static/uploads/_template.docx"
    doc_temp = docx.Document(doc_dir)
    tempreportdoc=screening(doc_temp)[1]
    return send_file(tempreportdoc, mimetype='application/msword',as_attachment=True, download_name='templatereport.docx')

@app.route("/highlightdoc")
def highlightdoc():
    if doc[1]:
        highlightdoc=screening(doc[1])[0]
    return send_file(highlightdoc, mimetype='application/msword',as_attachment=True, download_name='userdochighlight.docx')

@app.route("/reportdoc")
def reportdoc():
    if doc[1]:
        reportdoc=screening(doc[1])[1]
    return send_file(reportdoc, mimetype='application/msword',as_attachment=True, download_name='userdocreport.docx')

if __name__=="__main__":
    #app.run(debug=True)
    #app.run(host='0.0.0.0', port=8080, threaded=True)
    serve(app, host='0.0.0.0', port=8080, url_scheme="https")

