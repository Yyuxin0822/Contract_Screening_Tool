# -*- coding: utf-8 -*-
"""Copy of Main_Contract Language Risk Check.ipynb

Automatically generated by Colaboratory.

Original file is located at
    https://colab.research.google.com/drive/1EfoaXENT4eo2nOVWey7Sg3IJ8w8HD8yi

# Init

"""

"""**Import Libaries**"""


import re
import docx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from io import BytesIO

# # Create a new document and do some operations
# doc = Document()
# doc.add_paragraph("Hello, World!")

# # Save the document to a BytesIO stream
# stream = BytesIO()
# doc.save(stream)

# # Get the content of the document as a byte string
# document_content = stream.getvalue()

# # Close the stream
# stream.close()

# Now you have the document content in the `document_content` variable
"""**load doc and print info**

//It doesn't read tables, and pictures
"""

# doc_dir='./EXECUTED_1100009709_-_4100010177_-_Landscapes_of_the_Line__Forestation_Study_(Short_Form_Service_Agreement).docx'
# doc = docx.Document(doc_dir)

"""# Helper Functions and Class"""

def load_doc(doc):
    '''
    doc (Document object):
    Returns: a str of the document text
    '''
    text=f''
    for para in doc.paragraphs:
        text+='\n'+para.text
    return text

def find_sentencelist_with_keyword(keyword, text):
    pattern = r'\b\w*{}\w*\b'.format(re.escape(keyword))
    matches = re.finditer(pattern, text)
    output=[]
    for match in matches:
        start_index = match.start()
        end_index = match.end()

        # Find the start of the sentence
        sentence_start = text.rfind('.', 0, start_index) + 1

        # Find the end of the sentence
        sentence_end = text.find('.', end_index) + 1

        # Extract the sentence
        sentence = text[sentence_start:sentence_end].strip().replace("\n", " ")

        output.append(sentence)
    return output

class SensitiveWord:
  def __init__(self, text, warning,safeword="Your document looks safe on this topic.",asso=False,asso_text=[]):
        '''
        Constructs an instance of sensitiveword
        Inputs:
            count: Count of the appearance of the word in the doc
            warning: A warning str associated with the word
        '''
        self.text= text
        self.warning= warning
        self.safeword=safeword
        self.asso=asso
        self.asso_text=asso_text

  def get_text(self):
    '''
    Returns the sensitive word
    '''
    return self.text

  def get_count(self,doctext):
    '''
    Returns the count held by this sensitive word
    Inputs:the string to look up for the sensitive word
    '''
    return len(self.get_context(doctext))

  def get_warning(self):
    '''
    Returns the warning held by this sensitive word
    '''
    return self.warning

  def get_safeword(self):
    '''
    Returns the safeword held by this sensitive word
    '''
    return self.safeword

  def get_context(self,doctext):
    '''
    Returns the sentence list containing this sensitive word
    '''
    self.context=[]
    sentence=find_sentencelist_with_keyword(self.text, doctext)
    if self.asso:
      for s in sentence:
        for word in self.asso_text:
          if word in s:
            self.context.append(s)
    else:
      self.context=sentence
    return self.context

  def modify_doc(self,doc):
    '''
    Modify the original document to make the style changes for the keyword
    Inputs:
      doc: the original document object
      self.text: the text to make the style change
    '''
    keyword=self.text
    for paragraph in doc.paragraphs:
      if keyword in paragraph.text:
          for run in paragraph.runs:
              if keyword in run.text:
                if self.asso:
                  for word in self.asso_text:
                    if word in run.text:
                      run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                  run.font.highlight_color = WD_COLOR_INDEX.YELLOW


  def generate_keyword_report(self,doctext):
    '''
    Input:modifier(string)
    Return a string a string report based on get_count
    '''
    count=self.get_count(doctext)
    contextstring=""
    slist=self.get_context(doctext)
    for i in range(len(slist)):
      contextstring+=f'{(i+1)}. {slist[i]}\n\n'
    if self.asso:
      report= f"This document contains {count} sensitive word '{self.text}' in the context of '{self.asso_text}'.\n"
    else:
      report= f"This document contains {count} sensitive word '{self.text}'.\n"
    if count>0:
      report+=f'{self.get_warning()}\n'
      report+=f'They exist in these sentences.\n{contextstring}'
    else:
      report+=f'{self.get_safeword()}\n'
    return report







def screening(doc):

  doc_text=load_doc(doc)
  wordcount=len(doc_text.split())

  # print(doc_text)
  print("wordcount=:"+str(wordcount))

  # create a report document
  reportdoc = Document()

  worddict={}
  # create a dictionary that stores the associated information of the sensitive words
  # {sensitivewords: [warning,style]}
  reportdict={}
  # create a dictionary that stores the associated report information of the sensitive words
  # {sensitivewords: report sentence}

  # Add Disclamer
  disclaimer="This report doesn't examine any tables and picture contents.\n"
  disclaimer+=f"This report examined {wordcount} words."
  reportdoc.add_paragraph(disclaimer)




  def suggest_term(suggestion):
    return f'Suggested term on this topic:\n{suggestion}'

  def generate_section_report(sectiontitle=None,warningkeyword=None,warningsentence=None,suggestion=None,asso=False,asso_text=[]):
    '''
    Input:
      sectiontitle(str)
      warningkeyword(list)
      warningkeyword(str)
      suggestion(str)
      asso(boolean)
      asso_text(list)
    '''
    # Add Section Title
    if sectiontitle:
      reportdoc.add_heading(sectiontitle, level=3)


    # Raise warning
    if warningkeyword:
      for keyword in warningkeyword:
        worddict[keyword]=[]
        worddict[keyword].append(warningsentence)
        tempobject=SensitiveWord(keyword,worddict[keyword][0],asso=asso,asso_text=asso_text)
        tempobject.modify_doc(doc)
        reportdoc.add_paragraph(tempobject.generate_keyword_report(doc_text))


    # Suggest term
    if suggestion:
      reportdoc.add_paragraph(suggest_term(suggestion))

  """#  (NOCODE SECTION)Contract Language Examination

  ASHLEY HURD, CIC, CRM


  Professional Liability Insurance
  and Risk Management
  Presented by



  CONTRACTS 101
  Managing Risk Through Contract Language

  Topics of Discussion

  1.The Value of Professional Service Agreements

  2.Agreement Forms

  3.Incorporation by Reference of Prime Agreement

  4.Standard of Care

  5.Indemnity Clause

  6.Client Drafted Indemnities

  7.Code Compliance Issues

  8.Express Warranties and Guarantees

  9.Consequential Damages

  10.Assignments/Lender Requirements

  11.Dispute Avoidance/Dispute Resolution

  12.Claims Prevention Practices

  ---

  ---



  The Value of Professional Service Agreements

  ----Not Reflected in Code----

  An engineer is required by law to exercise a reasonable degree of care and skill when providing their services, even in the absence of a contract.
  The value and the benefit of the professional service agreement is that it clearly communicates proper roles and responsibilities.
  Contract negotiations with the client represent the prime opportunity to communicate with the client and the agreement should guide the relationship of the parties throughout the process.

  ---

  ---


  Agreement Forms

  ----Not Reflected in Code----

  The characteristics of the following agreement forms present special concerns:

  Client Generated Forms: should be examined closely to determine if they include an acceptable standard of care, indemnification and a clearly defined scope of services.

  Oral Agreements: while usually valid and binding, often result in an unclear understanding of the scope of services and inconsistent levels of expectations on the part of each party.

  ---

  Agreement Forms

  ----Not Reflected in Code----

  Purchase Orders: most are intended for product procurement and include product liability or express warranty provisions inappropriate for professional services. They also seldom identify the scope of professional services.

  Continuing Service Contracts: long-term service agreement can negate the application of a statute of repose, create a duty to advise a client of a change in codes and result in risk that exceeds the risk of the negotiated compensation.

  ---


  ---


  Incorporation by Reference of Prime Agreement

  ----Not Reflected in Code----

  Determine that the incorporated terms and conditions don’t create greater responsibility than the terms and conditions in the subcontract.

  At the end of the incorporation by reference sentence, insert the following:

  “…provided however that the Standard of Care and the Indemnification provisions set forth in this Subcontract take precedence over the Contract between Prime and its Client.”

  # Standard of Care

  ----Reflected in Code----
  1. Raise warning if *""perform to the highest standard of practice"", "highest standard", "highest"*
  2. Suggest "Services provided by the Engineer under this Agreement will be performed in a manner consistent with that degree of care and skill ordinarily exercised by members of the same profession currently practicing under similar conditions."

  Standard of Care
  Some clients will attempt in their contracts to change the standard of care language to require the firm to *perform to the highest standard of practice*, which is completely indefensible. Any language that seeks to raise the customary standard, places you in jeopardy of your professional liability coverage.

  If the Client drafts a contract clause that purports to raise the standard of care to a higher level, you must delete the offending words and revise to the standard level.


  Standard of Care Clause
  Services provided by the Engineer under this Agreement will be performed in a manner consistent with that degree of care and skill ordinarily exercised by members of the same profession currently practicing under similar conditions.
  """

  title="Standard of Care"
  keyword=['highest']
  warningsentence="Warning! Please examine if a contract clause that purports to raise the standard of care to a higher level"
  statement="Services provided by the Engineer under this Agreement will be performed in a manner consistent with that degree of care and skill ordinarily exercised by members of the same profession currently practicing under similar conditions."
  generate_section_report(title,keyword,warningsentence,statement)

  """# Indemnity Clause

  ----Reflected in Code----
  1.   Raise warning if *"defend"*

  Indemnity Clause

  A mutual indemnification:
    Consultant shall indemnify, and hold Client, its officers, and employees harmless from and against any and all claims, damages, liabilities, fines, penalties, losses, defense costs, including without limitation, reasonable attorneys' fees, and other liabilities (collectively “Losses”) to the extent caused by the negligence or willful misconduct of Consultant, its officers, agents, subcontractors or employees in connection with the project.

    Client shall indemnify, and hold Consultant, subcontractors, and employees harmless from and against any and all Losses to the extent caused by the negligence or willful misconduct of Client, its officers, agents, employees, contractors and agents in connection with the project.

    Neither Client nor Consultant shall be liable for any economic losses including without limitation, claims of loss of profits or any other indirect, incidental, or consequential damages.

  Indemnity Clause

  The essence of these words:
    Consultant shall indemnify and hold harmless Client from losses to the extent caused by the negligence or willful misconduct of Consultant in connection with the project.

  Do not agree to defend!
    Consultant shall indemnify, defend and hold harmless Client from losses caused by the negligence or willful misconduct of Consultant in connection with the project.

  Indemnity Clause

  While an engineer can agree to indemnify the client for the client’s reasonable costs of defense of a third-party claim once the negligence of the engineer is proved, assuming the defense of an allegation of negligence can be onerous. The firm risks its deductible on every defense tendered under the agreement and will find that it, and not its insurer, is paying the total cost of defense against an allegation.
  """

  title="Indemnity Clause"
  keyword=['defend']
  warningsentence="Warning! Please examine if a contract clause that agrees to defend"
  generate_section_report(title,keyword,warningsentence)

  """# Client Drafted Indemnities

  ----Reflected in Code----
  1. Suggest"This insurance does not apply to liability assumed by you under any contract; but this exclusion does not apply if you would have been liable in the absence of such contract."
  2. Raise warning if "in whole or in part" contained in a sentence that has "indemnities"

  3. Raise warning if "client’s partners, principles, officers, directors and employees"contained in a sentence that has "indemnities"
  client’s agents, contractors, attorneys, contract employees, lenders, volunteers or anyone else who is not directly part of the client entity.

  4. Suggest
  When all negotiation fails to delete the duty to defend, insist on one of  the following two statements to be added to the indemnity clause

    (1) At the preamble:
    To the fullest extent permitted by law, Engineer shall defend (to the extent covered by General Liability Insurance required by this Agreement) …

    (2) At the end of the indemnity:
    The Consultant’s obligation to defend professional negligence claims is limited to reimbursement of expenditure, including reasonable attorney fees and costs, incurred by an Indemnitee in defending claims or lawsuits, ultimately determined to be due to negligence acts or omissions of the Consultant or any of its employees or subconsultants.



  Client Drafted Indemnities

  If you sign a client’s indemnity that is not limited to your negligence, you are accepting liability beyond that is required by law and beyond that for which you are insured.
  Note: every professional liability policy has this exclusion:
  This insurance does not apply to liability assumed by you under any contract; but this exclusion does not apply if you would have been liable in the absence of such contract.


  Client Drafted Indemnities

  Often requests indemnity for claims caused in whole or in part by the architect, except for the sole negligence of the client.

  The inference is clear, the client intends that you pay all of the loss jointly caused by you and the client, even if the client is ninety-nine percent at fault and you are a mere one percent negligent. Keep in mind that even the ineptest lawyer can convince a jury  that the architect is one percent at fault for just about anything.

  Client Drafted Indemnities

  Often are uninsurable.
  Under the law, you have an obligation to perform services in a non-negligent manner. Your professional liability insurance covers you against damages resulting from your negligent professional acts, errors or omissions.
  If you sign a client’s indemnity that is not limited to your negligence, you are accepting liability beyond that required by law and that for which you are insured.
  Your professional liability policy includes the following exclusion: This insurance does not apply to liability assumed by you under any contract, but this exclusion does not apply if you would have been liable in the absence of such contract.


  Client Drafted Indemnities

  Attempt to include inappropriate parties as indemnities.
  It is not unusual to indemnify a client’s partners, principles, officers, directors and employees.
  You should not agree to indemnify a client’s agents, contractors, attorneys, contract employees, lenders, volunteers or anyone else who is not directly part of the client entity.
  You do not owe these parties the same obligations, and they can always seek their own legal remedies should your negligence somehow damage them.


  When all negotiation fails to delete the duty to defend, insist on one of  the following two statements to be added to the indemnity clause.

  (1) At the preamble:
    To the fullest extent permitted by law, Engineer shall defend (to the extent covered by General Liability Insurance required by this Agreement) …
  (2) At the end of the indemnity:
    The Consultant’s obligation to defend professional negligence claims is limited to reimbursement of expenditure, including reasonable attorney fees and costs, incurred by an Indemnitee in defending claims or lawsuits, ultimately determined to be due to negligence acts or omissions of the Consultant or any of its employees or subconsultants.
  """

  # Add Section Title
  title="Client Drafted Indemnities"
  reportdoc.add_heading(title, level=3)

  # This part works with asso_sentence

  warningword=['in whole or in part']
  warningsentence="Warning! Please examine if a contract clause that requests indemnity for claims caused in whole or in part by the architect, except for the sole negligence of the client."
  # Raise warning
  for keyword in warningword:
      worddict[keyword]=[]
      worddict[keyword].append(warningsentence)
      tempobject=SensitiveWord(keyword,worddict[keyword][0],asso=True,asso_text=["indemnity"])
      tempobject.modify_doc(doc)
      reportdoc.add_paragraph(tempobject.generate_keyword_report(doc_text))

  # This part works with asso_sentence

  warningword=['partners', 'principles', 'officers', 'directors', 'employees','agents', 'contractors', 'attorneys','contract employees', 'lenders', 'volunteers' or 'anyone' ]
  warningsentence="Warning! Please examine if a contract clause that attempts to include inappropriate parties as indemnities."
  # Raise warning
  for keyword in warningword:
      worddict[keyword]=[]
      worddict[keyword].append(warningsentence)
      tempobject=SensitiveWord(keyword,worddict[keyword][0],asso=True,asso_text=["indemnity"])
      tempobject.modify_doc(doc)
      reportdoc.add_paragraph(tempobject.generate_keyword_report(doc_text))

  suggestion="This insurance does not apply to liability assumed by you under any contract; but this exclusion does not apply if you would have been liable in the absence of such contract."
  suggestion+="When all negotiation fails to delete the duty to defend, insist on one of  the following two statements to be added to the indemnity clause.\n (1) At the preamble: \n To the fullest extent permitted by law, Engineer shall defend (to the extent covered by General Liability Insurance required by this Agreement) …\n (2) At the end of the indemnity: \n The Consultant’s obligation to defend professional negligence claims is limited to reimbursement of expenditure, including reasonable attorney fees and costs, incurred by an Indemnitee in defending claims or lawsuits, ultimately determined to be due to negligence acts or omissions of the Consultant or any of its employees or subconsultants."
  generate_section_report(None,None,None,suggestion)

  """# Code Compliance Issues

  ----Reflected in Code----

  1. Raise warning if "laws" "codes" appear in the context of "comply", and "complaince"
  2.  "Fix it by committing only to exercise the standard of care to comply."

    "The consultant shall put forth reasonable professional efforts to comply with codes, regulations, laws…."

    "The consultant shall exercise usual and customary professional care in their efforts to comply with all codes, regulations, laws that are in effect at the time services are rendered."

  3.   Search “Owner/Client recognizes that some requirements and codes are subject to multiple reasonable interpretations and that the Consultant will exercise the Standard of Care to interpret and apply the same properly for the project.”



  What if your contract says …
  Consultant’s services shall comply with all codes, regulations, and ordinances applicable to the Project.

  Difficult to comply!

  Thousands of laws, codes and regulations that relate to construction are on the books; all are subject to change, and some are open to interpretation.
  It is not unusual to find that a given regulation may conflict with another.
  If you can only adhere to one – and you have agreed to comply with all….


  Compliance with Laws

  Don’t warrant compliance with laws and codes.

  When you see a clause stating that the consultant shall comply with all laws, codes, etc., this creates a warranty. Fix it by committing only to exercise the standard of care to comply.

  Alternatives
  The consultant shall put forth reasonable professional efforts to comply with codes, regulations, laws….

  The consultant shall exercise usual and customary professional care in their efforts to comply with all codes, regulations, laws that are in effect at the time services are rendered.


  A Compromise on Code Compliance

  A final compromise: Add this sentence if you can’t come to an agreement to add Standard of Care generally:

  “Owner/Client recognizes that some requirements and codes are subject to multiple reasonable interpretations and that the Consultant will exercise the Standard of Care to interpret and apply the same properly for the project.”
  """

  title="Code Compliance Issue"
  keyword=['law','code','regulation','ordinance']
  warningsentence="Warning! Please examine if a contract clause that consultant’s services shall comply with all codes, regulations, and ordinances applicable to the project"
  statement="Fix it by committing only to exercise the standard of care to comply.\nThe consultant shall put forth reasonable professional efforts to comply with codes, regulations, laws….\n The consultant shall exercise usual and customary professional care in their efforts to comply with all codes, regulations, laws that are in effect at the time services are rendered. \n Owner/Client recognizes that some requirements and codes are subject to multiple reasonable interpretations and that the Consultant will exercise the Standard of Care to interpret and apply the same properly for the project."
  generate_section_report(title,keyword,warningsentence,statement,asso=True,asso_text=["comply","compliance"])

  """# Express Warranties and Guarantees

  ----Reflected in Code----

  1.   Raise warning if  “warrant” or “guarantee” “ensure” or “confirm” Words such as “insure,” “ensure,” or “assure” are indicative of an express warranty or guarantee. Similarly, words such as “complete” or “every” and phrases such as “fit for the intended purpose” can create warranties
  2.  What is the sensitive word for this? "Consultant shall perform its obligations under this Agreement in the most efficient and economical manner consistent with the Owner’s best interests."


  Express Warranties and Guarantees

  Express warranties and guarantees establish liability even though no proof of negligence is required.
  Courts have not extended the duty to provide a guarantee to engineering professionals because they provide services based on judgement and expertise; a design professional applies its professional skills and reasoning to a unique set of conditions for each project.
  Professional liability policies accordingly exclude coverage for claims arising out of express warranties or guarantees.

  Express Warranties and Guarantees

  The words “warrant” or “guarantee” mean to “ensure” or “confirm” that a standard has been met absolutely.

  Words such as “insure,” “ensure,” or “assure” are indicative of an express warranty or guarantee. Similarly, words such as “complete” or “every” and phrases such as “fit for the intended purpose” can create warranties.

  Warranties & Guarantees

  Examples of a scope with an inappropriate warranty:

  Develop a schedule to assure timely completion.

  Perform other services to insure the accuracy of cost and schedule commitments.

  Consultant shall perform its obligations under this Agreement in the most efficient and economical manner consistent with the Owner’s best interests.
  """

  title="Express Warranties and Guarantees"
  keyword=['warrant','guarantee','ensure','confirm','insure','ensure','assure','complete','every','fit for the intended purpose']
  warningsentence="Warning! Please examine if indicative of an express warranty or guarantee"
  generate_section_report(title,keyword,warningsentence)

  """# Consequential Damages

  ----Reflected in Code----

  1.   Raise warning if "Consequential Damages"
  2.   Suggest Terms"
  Notwithstanding any other provision of this Agreement, neither party shall be liable to the other for any consequential damages incurred due to the fault of the other party, regardless of the nature of this fault or whether it was committed by the Client or the Design Professional, their employees, agents, subconsultants or subcontractors. Consequential damages include, but are not limited to, loss of use and loss of profit.
  "

  Consequential Damages

  The Problem:
  If you are to be held responsible for consequential damages, you could be sued for damages wholly out of proportion to the cost of repairing the actual damage.


  The Solution:
  Add a provision that makes certain neither you nor your client will be held responsible for consequential damages.

  Consequential Damages

  Notwithstanding any other provision of this Agreement, neither party shall be liable to the other for any consequential damages incurred due to the fault of the other party, regardless of the nature of this fault or whether it was committed by the Client or the Design Professional, their employees, agents, subconsultants or subcontractors. Consequential damages include, but are not limited to, loss of use and loss of profit.
  """

  title="Consequential Damages"
  keyword=['consequential damages']
  warningsentence="Warning! Please examine if you are to be held responsible for consequential damages"
  statement="Notwithstanding any other provision of this Agreement, neither party shall be liable to the other for any consequential damages incurred due to the fault of the other party, regardless of the nature of this fault or whether it was committed by the Client or the Design Professional, their employees, agents, subconsultants or subcontractors. Consequential damages include, but are not limited to, loss of use and loss of profit."
  generate_section_report(title,keyword,warningsentence,statement)

  """# Assignments/Lender Requirements

  ----Reflected in Code----
  1.   Raise warning if "lenders"
  2. Suggest Terms"
  Neither party may assign, sublet or transfer any rights under or interest in this Agreement without the written consent of the other. Unless specifically stated to the contrary in any written consent to assignment, no assignment will release or discharge the assignor from any duty or responsibility under this Agreement.
  "

  Assignments/Lender Requirements

  Issue: Your client insists on assigning its contractual rights to a lender that requires your consent to the assignment and your certification of project information.
  Concerns: From a professional liability perspective, is the architect extending its liability through its statements to the lender. Does the assignment contain express warranty or guarantee language. Are you comfortable with providing services to an unknown client and be properly compensated.

  Assignments/Lender Requirements

  Additional Concerns:
  Is your new client willing to assume the obligations of the original client?
  Is there any recourse for the payment of uncollected fees and future fees, including any costs generated by the assignment?
  Do you want to provide services to the lender or to a subsequent assignee?
  Is the normal legal liability being extended in time or in scope?

  Assignments/Lender Requirements

  The Solution:
  Professional liability exposures and uninsurable risks can increase significantly because of the language in a consent to assignment. The best protection against such risk is to include the following language in your agreement:

  Neither party may assign, sublet or transfer any rights under or interest in this Agreement without the written consent of the other. Unless specifically stated to the contrary in any written consent to assignment, no assignment will release or discharge the assignor from any duty or responsibility under this Agreement.
  """

  title="Assignments/Lender Requirements"
  keyword=['lenders']
  warningsentence="Warning! Please examine if your client insists on assigning its contractual rights to a lender"
  statement="Neither party may assign, sublet or transfer any rights under or interest in this Agreement without the written consent of the other. Unless specifically stated to the contrary in any written consent to assignment, no assignment will release or discharge the assignor from any duty or responsibility under this Agreement. "
  generate_section_report(title,keyword,warningsentence,statement)

  """#  (NOCODE SECTION)Contract Language Examination

  ---



  ---


  Dispute Avoidance/Dispute Resolution

  ----Not Reflected in Code----

  Glossary:

  Mediation: Facilitated negotiation. No resolution absent agreement.

  Arbitration: Private decision maker.

  Litigation: Public litigation in state of federal courts.

  Dispute Avoidance/Dispute Resolution
  Problems with the Arbitration Option

  Arbitration tends to be more favorable to contractors than design professionals.
  Reasons?
  Contractors have more claims and hire arbitrators more often.
  If you agree to arbitrate, it is very important to select an arbitrator with design professional experience.
  Can be slow and expensive.

  Dispute Avoidance/Dispute Resolution
  Benefits of Litigation

  Motions to Dismiss
  Motion for Summary Judgement
  Courts enforce terms of contract and grant appropriate motions.
  Decision is clearer than an arbitration decision concerning findings of fact and conclusions of law.
  Appealable.

  ---



  ---


  Claims Prevention Practices

  ----Not Reflected in Code----

  Documentation
  Consistent documentation increases your credibility.
  Documentation includes written records, audio/video tapes, photographs and computer records.
  Observations should be promptly recorded.
  Record observed facts – avoid unneeded editorial comments.
  Assume all project records are discoverable.
  Record retention should at least track state statues plus repose.


  Claims Prevention Practices

  Records Preservation
  Contracts and amendments.
  Meeting notes & phone logs.
  Advisory letters.
  Invoices.
  Product information & manufacturer guarantees.
  Site visit reports & checklists.
  Minutes of construction site meetings and field observations.
  Project completion documentation.
  Photographs.

  Claims Prevention Practices

  Memos of formal and informal telephone conversations and conferences.
  Written authorization to proceed at each phase.
  Copies of owner furnished data.
  Documentation of engineering recommendations and the owner’s decision and response thereto.
  Owners’ acceptance of the work as substantially complete.
  Minutes of construction site meetings and field observations.
  Notification of estimated cost increases with owner’s response thereto.


  Claims Prevention Practices

  Write it Right!
  All correspondence can create liability exposures, including proposals, agreements, letters, memos, notes and emails.
  Abide by the following rules to avoid liability:
  Keep it factual, no conjecture.
  Use plain language, avoid technical jargon.
  Avoid editorial comments such as:
  We should have …
  Our design is defective …
  We probably have some liability …
  We will have to pay to correct it …
  Never state an opinion on responsibility.


  Claims Prevention Practices

  Write it Right!
  Examples of actual documents:
  Many of these design goofs I’ve mentioned would be readily apparent to any reasonably knowledgeable observer.  I would not recommend taking a prospective client to this plant.

  We were in error in not picking up the detailer’s mistake.  So certainly we must share the responsibility for the fact that the reinforcing was improperly detailed and fabricated.

  The procedure for shop drawing review needs to be formalized.  We have held drawings longer than necessary through failure to follow standard procedure.

  ---

  ---

  Legal Disclaimer
  This presentation contains suggested edits to contract
  language. These are intended to assist our clients to:

  Evaluate compliance with certain insurance requirements;
  Identify contract requirements that may exceed coverage under insurance policies; and
  Suggest areas of the contract that may pose increased risk to the client.

  The information contained in this presentation is not legal advice. This should not replace consultation with the firm’s legal counsel.


  ASHLEY HURD, CIC, CRM


  Questions & Answers
  Presented by

  # Generate Report and Output
  """

  # Save the document to a BytesIO stream
  stream_highlight = BytesIO()
  doc.save(stream_highlight)
  stream_highlight.seek(0)

  stream_report = BytesIO()
  reportdoc.save(stream_report)
  stream_report.seek(0)
  
  return (stream_highlight,stream_report)



if __name__=="__main__":
  pass
  #app.run(debug=True)
  #app.run(host='0.0.0.0', port=8080, threaded=True)
  #serve(app, host='0.0.0.0', port=8080, url_scheme="https")
